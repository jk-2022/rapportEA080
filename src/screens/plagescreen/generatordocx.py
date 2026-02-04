import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from flet import *


class GeneratorDOCX():
    def __init__(self, donnees: list, titre: str = "Rappor", filename: str = "rapport.docx"):
        """
        :param donnees: Liste de dictionnaires contenant les données (avec clés 'date', 'titre', 'contenu')
        :param titre: Titre principal du document
        :param filename: Nom du fichier DOCX à générer
        """
        self.donnees = donnees
        self.titre = titre.upper()
        self.filename = filename

    def set_col_width(self, cell, width_inches):
        cell.width = Inches(width_inches)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)

    def create_docx(self):
        doc = Document()

        # Titre
        titre = doc.add_heading(self.titre, level=1)
        titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph("")

        # Table: 3 colonnes (Date, Titre, Contenu)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False  # ❗ important : empêche l'ajustement automatique

        # Largeurs : total ≤ 6.0 pouces (format A4 utile)
        col_widths = [1.0, 1.5, 3.5]

        headers = ["Date", "Titre", "Contenu"]
        hdr_cells = table.rows[0].cells

        for i, text in enumerate(headers):
            para = hdr_cells[i].paragraphs[0]
            run = para.add_run(text)
            run.bold = True
            self.set_col_width(hdr_cells[i], col_widths[i])
        # print(self.donnees)
        for ligne in self.donnees:
            content=json.loads(ligne['content'])
            info=content.pop("info_ouvrage")
            type_ouvrage=content.pop("type_ouvrage")
            text_info=""
            text_cont=""
            if type_ouvrage:
                text_info+=f">Type Ouvrage : {type_ouvrage}\n"
                for key,value in info.items():
                    if value=="":
                        text_info=text_info
                    else:
                        text_info +=f"{key} : {value}   "
                

            for key,value in content.items():
                texte=Text(f"> {key.upper()}\n {value}")
                text_cont += f"{texte.value}\n"
            
            text_final=""
            if text_info!="":
                text_final= f"{text_info}\n{text_cont}"
            else:
                text_final=text_cont

            row_cells = table.add_row().cells
            row_cells[0].text = ligne.get("rapport_date", "")
            row_cells[1].text = ligne.get("title", "")
            row_cells[2].text = text_final
            for i, width in enumerate(col_widths):
                self.set_col_width(row_cells[i], width)

        doc.save(self.filename)

