import os
from flet import *
from fpdf import FPDF
from fpdf.enums import XPos, YPos

from fpdf.enums import XPos, YPos

from docx import Document
import csv
import sqlite3
import json

from allpath import AllPath

path=AllPath()
path_data=path.path_data()
generated_docs=path.path_generated_docs()

DB_PATH=os.path.join(path_data,"rapport.db")

from screens.rapportscreen.rapportupdateform import RapportUpdateForm

def convertir_dict_to_text(content):
    content=json.loads(content)
    info=content.pop("info_ouvrage")
    type_ouvrage=content.pop("type_ouvrage")
    text_info=""
    text_cont=""
    if type_ouvrage:
        text_info+=f"> Type Ouvrage : {type_ouvrage}\n"
        for key,value in info.items():
            if value=="":
                text_info=text_info
            else:
                text_info +=f"{key} : {value}   "
        

    for key,value in content.items():
        texte=Text(f"> {key.upper()}\n {value}")
        text_cont += f"{texte.value}\n"
    
    text_final=""
    if text_info!="":
        text_final= f"{text_info}\n{text_cont}"
    else:
        text_final=text_cont

    return text_final


class RapportCard(Card):
    def __init__(self, page: Page, rapport, formcontrol):
        super().__init__()
        self.page=page
        self.expand=True
        self.elevation=10
        self.rapport=rapport
        self.formcontrol=formcontrol
        self.contenue_cont=Column(
                            controls=[
                                Row(
                                    [
                                        Text(f"{self.rapport['title']}")
                                    ],alignment=MainAxisAlignment.CENTER
                                    ),
                                Row([
                                    Text(f"Rapport du {rapport['rapport_date']}")
                                ],alignment=MainAxisAlignment.CENTER
                                ), 
                            ],expand=True
                            )
        
                            
            
        self.content=Container(
            # on_click=lambda e: self.selectprojet(e),
            padding=padding.all(10),
            data=rapport,
            ink=True,
            expand=True,
            content=Container(
                content=Column(
                    [
                        self.contenue_cont,
                        Row([
                            IconButton(icon=Icons.EDIT, on_click= self.go_update_page),
                            IconButton(icon=Icons.DELETE, on_click= self.show_delete_rapport),
                            PopupMenuButton(items=[
                                PopupMenuItem(text="Générer PDF", on_click=self.showGenerate_pdf),
                                PopupMenuItem(text="Générer DOCX", on_click=self.showGenerate_docx),
                                PopupMenuItem(text="Générer CSV", on_click=self.showGenerate_csv),
                                ])
                            ],alignment=MainAxisAlignment.END
                            )
                        ]
                    ),
                padding=10,
                expand=True
            )
        )
        self.convertir_bd_content(rapport['content'])

    def convertir_bd_content(self, cont):
        content=json.loads(cont)
        info=content.pop("info_ouvrage")
        type_ouvrage=content.pop("type_ouvrage")
        
        if type_ouvrage:
            text_info=""
            for key,value in info.items():
                if value=="":
                    text_info=text_info
                else:
                    text_info +=f"{key} : {value}   "
            if text_info!="":
                text=Text(f">> Type Ouvrage : {type_ouvrage}\n",size=13,spans=[TextSpan(f"{text_info}",
                                        TextStyle( size=11,italic=True, weight=FontWeight.W_200)
                                        )
                                        ])
                self.contenue_cont.controls.append(text)

        for key,value in content.items():
            texte=Text(f">> {key.upper()}\n",size=13,spans=[TextSpan(f"{value}",
                                       TextStyle( size=11,italic=True, weight=FontWeight.W_200)
                                       )
                                    ])
            self.contenue_cont.controls.append(texte)
        

    def go_update_page(self,e):
        self.page.client_storage.set('rapport',self.rapport)
        self.page.go("/update-rapport")

    def show_delete_rapport(self,e):
        title=self.rapport['title']
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Confirmation"),
            content=Text(f"Voulez-vous supprimer {title} ?"),
            actions=[
                TextButton("Non", on_click=self.close_dlg),
                TextButton("Oui", on_click=self.del_rapport),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
            content_padding=0
        )
        self.page.open(self.dlg_modal)
        self.page.update()
    
    def showGenerate_pdf(self,e):
        title=self.rapport['title']
        titlefield=TextField(value=title, expand=True,multiline=True)
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Nom du fichier"),
            content=titlefield,
            actions=[
                TextButton("Annuler", on_click=self.close_dlg),
                TextButton("Exporter", on_click = lambda e : self.generate_pdf(titlefield.value)),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
        )
        self.page.open(self.dlg_modal)
        self.page.update()

    def showGenerate_csv(self,e):
        title=self.rapport['title']
        titlefield=TextField(value=title, expand=True,multiline=True)
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Nom du fichier"),
            content=titlefield,
            actions=[
                TextButton("Annuler", on_click=self.close_dlg),
                TextButton("Exporter", on_click = lambda e : self.generate_csv(titlefield.value)),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
        )
        self.page.open(self.dlg_modal)
        self.page.update()

    def showGenerate_docx(self,e):
        title=self.rapport['title']
        titlefield=TextField(value=title, expand=True,multiline=True)
        self.dlg_modal = AlertDialog(
            modal=True,
            title=Text("Nom du fichier"),
            content=titlefield,
            actions=[
                TextButton("Annuler", on_click=self.close_dlg),
                TextButton("Exporter", on_click = lambda e : self.generate_docx(titlefield.value)),
            ],
            actions_alignment= MainAxisAlignment.END,
            on_dismiss=lambda e: print("Modal dialog dismissed!"),
        )
        self.page.open(self.dlg_modal)
        self.page.update()
        

    def generate_docx(self, title):
        report_id=self.rapport["rid"]
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT title, content FROM rapports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()
        if row:
            title, content = row
            doc = Document()
            doc.add_heading(title, level=1)
            text_final=convertir_dict_to_text(content)
            doc.add_paragraph(text_final)
            doc.save(f"{generated_docs}/rapport_{title}.docx")
        self.close_dlg(e=None)
            
    def generate_csv(self, title):
        report_id=self.rapport["rid"]
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT title, content, rapport_date FROM rapports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()
        if row:
            text_final=convertir_dict_to_text(row[1])
            row=list(row)
            row[1]=str(text_final)
            with open(f"{generated_docs}/rapport_{title}.csv", mode="w", newline="", encoding="utf-8") as file:
                writer = csv.writer(file,delimiter=";")
                writer.writerow(["Titre", "Contenu", "Date"])
                writer.writerow(row)
        self.close_dlg(e=None)

    def generate_pdf(self, title):
        report_id=self.rapport["rid"]
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT title, content, rapport_date FROM rapports WHERE id = ?", (report_id,))
        row = cursor.fetchone()
        conn.close()

        if row:
            title, content, rapport_date = row
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("helvetica", "B", 16)

            # Titre du rapport
            # pdf.cell(0, 10, title, ln=True, align='C')
            pdf.cell(0, 10, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

            # Semaine
            pdf.set_font("helvetica", "", 12)
            pdf.cell(0, 10, f"Rapport du {rapport_date}", new_x=XPos.LMARGIN, new_y=YPos.NEXT, align='C')

            pdf.ln(10)

            # Contenu
            text_final=convertir_dict_to_text(content)
            pdf.multi_cell(0, 10, text_final)

            # Sauvegarde
            file_path = f"{generated_docs}/rapport_{title}.pdf"
            pdf.output(file_path)

            self.page.open(SnackBar(Text(f"rapport_{title} est exporter avec succès")))
        self.close_dlg(e=None)

        
    def close_dlg(self,e):
        try:
            self.page.close(self.dlg_modal)
            self.page.update()
        except:
            pass
        
    def del_rapport(self,e):
        rid=int(self.rapport['rid'])
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM rapports WHERE id=?", (rid,))
        conn.commit()
        conn.close()
        self.page.close(self.dlg_modal)
        self.formcontrol.load_rapports()
        
    