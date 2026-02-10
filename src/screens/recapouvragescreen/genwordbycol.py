import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from mystorage import get_value 

archive_path=get_value("archive_path")

def generer_pv_word_auto(
    titre: str,
    sous_titre: str,
    donnees: dict,
    nom_fichier: str = "proces_verbal.docx"
):
    doc = Document()

    # ===== Titre =====
    title = doc.add_heading(titre, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(sous_titre)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    doc.add_paragraph(
        "L’an deux mille vingt-six (2026), il a été procédé à la réception "
        "provisoire des travaux conformément aux données techniques ci-dessous."
    )

    # ===== 1️⃣ Filtrage des sections non vides =====
    sections_valides = {}

    for nom_section, contenu in donnees.items():
        if isinstance(contenu, list) and len(contenu) > 0:
            sections_valides[nom_section] = contenu[0]  # premier dict
        elif isinstance(contenu, dict) and contenu:
            sections_valides[nom_section] = contenu

    if not sections_valides:
        doc.add_paragraph("Aucune donnée technique disponible.")
        doc.save(nom_fichier)
        return nom_fichier

    # ===== 2️⃣ Tableau conteneur dynamique =====
    nb_sections = len(sections_valides)
    container = doc.add_table(rows=1, cols=nb_sections)
    container.style = "Table Grid"

    # ===== 3️⃣ Remplissage =====
    for col_index, (section, valeurs) in enumerate(sections_valides.items()):
        cell = container.rows[0].cells[col_index]

        # Titre de section
        p = cell.paragraphs[0]
        p.text = section.upper()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sous-tableau 2 colonnes
        table = cell.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        for cle, valeur in valeurs.items():
            if valeur is None or valeur == "":
                valeur = "-"
            row = table.add_row().cells
            row[0].text = str(cle)
            row[1].text = str(valeur)

    # ===== Décision =====
    doc.add_paragraph("\n")
    doc.add_heading("Décision", level=2)
    doc.add_paragraph(
        "Au vu des éléments disponibles, l’ouvrage est déclaré conforme "
        "et réceptionné provisoirement."
    )

    # ===== Signatures =====
    doc.add_heading("Signatures", level=2)
    table_sign = doc.add_table(rows=1, cols=3)
    table_sign.style = "Table Grid"

    headers = table_sign.rows[0].cells
    headers[0].text = "Qualité"
    headers[1].text = "Nom et Prénoms"
    headers[2].text = "Signature"

    roles = [
        "Représentant du Projet",
        "Représentant de l’Entreprise",
        "Représentant de la Collectivité",
    ]

    for role in roles:
        row = table_sign.add_row().cells
        row[0].text = role
        row[1].text = ""
        row[2].text = ""

    path_fichier=os.path.join(archive_path,nom_fichier)
    doc.save(path_fichier)

    return path_fichier
