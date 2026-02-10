import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from mystorage import get_value 

archive_path=get_value("archive_path")

def generer_proces_verbal_word(
    titre: str,
    sous_titre: str,
    donnees: dict,
    nom_fichier: str = "proces_verbal.docx"
    ):
    """
    Génère un procès-verbal Word (.docx) à partir de données structurées
    """

    doc = Document()

    # ===== Titre =====
    title = doc.add_heading(titre, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(sous_titre)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    # ===== Texte juridique standard =====
    doc.add_paragraph(
        "L’an deux mille vingt-six (2026), il a été procédé à la réception "
        "provisoire des travaux conformément aux données techniques ci-dessous."
    )

    # ===== Génération dynamique des tableaux =====
    section_num = 1
    for section, valeurs in donnees.items():
        # print(valeurs)
        if donnees[section]:
            doc.add_heading(f"{section_num}. {section.capitalize()}", level=2)

            table = doc.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            for cle, valeur in valeurs[0].items():
                row = table.add_row().cells
                row[0].text = str(cle)
                row[1].text = str(valeur)

            section_num += 1

    # ===== Conclusion =====
    # doc.add_heading(f"{section_num}. Décision", level=2)
    # doc.add_paragraph(
    #     "Au vu des résultats obtenus, l’ouvrage est déclaré conforme "
    #     "et réceptionné provisoirement."
    # )

    # ===== Signatures =====
    # doc.add_heading(f"{section_num + 1}. Signatures", level=2)

    # table_sign = doc.add_table(rows=1, cols=3)
    # table_sign.style = "Table Grid"

    # headers = table_sign.rows[0].cells
    # headers[0].text = "Qualité"
    # headers[1].text = "Nom et Prénoms"
    # headers[2].text = "Signature"

    # roles = [
    #     "Représentant du Projet",
    #     "Représentant de l’Entreprise",
    #     "Représentant de la Collectivité",
    # ]

    # for role in roles:
    #     row = table_sign.add_row().cells
    #     row[0].text = role
    #     row[1].text = ""
    #     row[2].text = ""

    # ===== Sauvegarde =====
    path_fichier=os.path.join(archive_path,nom_fichier)
    doc.save(path_fichier)

    return path_fichier
