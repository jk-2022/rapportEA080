import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from mystorage import get_value 
archive_path=get_value("archive_path")

def generer_pv_word_final(
    titre: str,
    sous_titre: str,
    donnees: dict,
    nom_fichier: str = "proces_verbal.docx"
):
    """
    Génère un procès-verbal Word (.docx) avec :
    - marges réduites
    - police 10 pt pour les tableaux
    - sections dynamiques (Ouvrage / Foration / Pompage)
    - tableaux sur une seule ligne si possible
    """

    doc = Document()

    # ===== Réduction des marges =====
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # ===== Titre et sous-titre =====
    title = doc.add_heading(titre, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(sous_titre)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n")

    # ===== Texte juridique =====
    doc.add_paragraph(
        "L’an deux mille vingt-six (2026), il a été procédé à la réception "
        "provisoire des travaux conformément aux données techniques ci-dessous."
    )

    # ===== Filtrage des sections non vides =====
    sections_valides = {}

    for nom_section, contenu in donnees.items():
        if isinstance(contenu, list) and len(contenu) > 0:
            sections_valides[nom_section] = contenu[0]  # premier dict
        elif isinstance(contenu, dict) and contenu:
            sections_valides[nom_section] = contenu

    if not sections_valides:
        doc.add_paragraph("Aucune donnée technique disponible.")
        doc.save(nom_fichier)
        return nom_fichier

    # ===== Tableau conteneur dynamique (1 ligne, N colonnes) =====
    nb_sections = len(sections_valides)
    container = doc.add_table(rows=1, cols=nb_sections)
    container.style = "Table Grid"

    # ===== Remplissage des tableaux =====
    for col_index, (section_nom, valeurs) in enumerate(sections_valides.items()):
        cell = container.rows[0].cells[col_index]

        # Titre de la section
        p = cell.paragraphs[0]
        p.text = section_nom.upper()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Sous-tableau 2 colonnes
        table = cell.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        
        cle_ignored=["id","projet_id","ouvrage_id","created_at"]
        for cle, valeur in valeurs.items():
            if cle in cle_ignored or valeur=="":
            # if valeur is None or valeur == "":
            #     valeur = "-"
                pass
            else:
                row = table.add_row().cells
                row[0].text = str(cle)
                row[1].text = str(valeur)

                # Police 10 pt pour chaque cellule
                for cell_inner in row:
                    for paragraph in cell_inner.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

    # # ===== Décision =====
    # doc.add_paragraph("\n")
    # doc.add_heading("Décision", level=2)
    # doc.add_paragraph(
    #     "Au vu des éléments disponibles, l’ouvrage est déclaré conforme "
    #     "et réceptionné provisoirement."
    # )

    # # ===== Signatures =====
    # doc.add_heading("Signatures", level=2)
    # table_sign = doc.add_table(rows=1, cols=3)
    # table_sign.style = "Table Grid"

    # headers = table_sign.rows[0].cells
    # headers[0].text = "Qualité"
    # headers[1].text = "Nom et Prénoms"
    # headers[2].text = "Signature"

    # roles = [
    #     "Représentant du Projet",
    #     "Représentant de l’Entreprise",
    #     "Représentant de la Collectivité",
    # ]

    # for role in roles:
    #     row = table_sign.add_row().cells
    #     row[0].text = role
    #     row[1].text = ""
    #     row[2].text = ""

    # ===== Sauvegarde =====
    path_fichier=os.path.join(archive_path,nom_fichier)
    doc.save(path_fichier)

    return path_fichier
